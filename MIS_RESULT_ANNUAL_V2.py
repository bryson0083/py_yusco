# -*- coding: utf-8 -*-
"""
年度工作實績表文件檔產生 V2

@author: Bryson Xue

@Note: 
	1. 讀取 SQL SERVER 整年度資料工作實績表資料
	2. 執行本程式產生WORD文件檔案
	3. 直接讀取資料庫，轉word文件

@Ref:

"""
import sys
import pyodbc 
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
import datetime
from dateutil import parser
from dateutil.parser import parse

err_flag = False

#Get current date
str_date = str(datetime.datetime.now())
yyyy = parser.parse(str_date).strftime("%Y")
yyy = int(yyyy) - 1911

#設定抓取日期區間
dt_st = yyyy + '0101'
dt_ed = yyyy + '1231'

#設定抓取日期區間(手動)
#dt_st = '20171120'
#dt_ed = '20171120'

#建立資料庫連線
conn = pyodbc.connect("Driver={SQL Server Native Client 10.0};"
					  "Server=ntsr12;"
					  "Database=per;"
					  "uid=sqluser;pwd=sqluser")

cursor = conn.cursor()

#讀取工作實績資料，以系統代號、事件日期排序
try:
	sql_str = "select * from ( select sys_no, title, useful, event_date, emplno, ( select username from secuser where userno = 'yu' + emplno) as username, ( select deptid from yd2a110m where emplno = yd2a600m.emplno) as deptid from yd2a600m where event_date >= '" + dt_st + "' and event_date <= '" + dt_ed + "' ) as a where a.deptid like 'F23%' order by sys_no, event_date"
	cursor.execute(sql_str)
	rt_list = cursor.fetchall()
	#print(rt_list)

except Exception as e:
	err_flag = True
	print("EXEC SQL:\n" + sql_str + "\n")
	print("Err:\n" + str(e.args))
	sys.exit("SQL Error 程式異常終止...")

labels = ['sys_no', 'title', 'useful', 'event_date', 'emplno', 'username', 'deptid']
df = pd.DataFrame.from_records(rt_list, columns=labels)

#以下開始產生WORD檔程序
#取df的資料筆數
row_cnt = df.shape[0]
#print(str(df.shape[0]))

#取df的資料欄位數
col_cnt = df.shape[1]
#print(str(df.shape[1]))

document = Document()

#文件格式風格設定
obj_styles = document.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(16)
#obj_font.name = 'Times New Roman'
obj_font.name = '標楷體'
obj_font.color.rgb = RGBColor(0x0, 0x0, 0xff)

#文件內文段落
paragraph = document.add_paragraph()
paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

#文件抬頭
run = paragraph.add_run(str(yyy) + '年度生產資訊處工作實績一覽表', style = 'CommentsStyle')

#加入表格
t = document.add_table(rows = row_cnt+1, cols = col_cnt - 3, style='Light List')
#t.style = 'Table Grid'
#t.style = 'Medium Shading 2 Accent 1'
#t.style = 'Medium List 2'
#t.style = 'Light List'

# add the header rows.
for j in range(df.shape[-1] - 3):
	t.cell(0,j).text = df.columns[j]

# add the rest of the data frame
prev_sysid = ''
curr_sysid = ''
for i in range(df.shape[0]):
	curr_sysid = str(df.values[i,0])
	#print(curr_sysid)

	diff_flag = False
	if i == 0:
		prev_sysid = curr_sysid
		diff_flag = True
		num = 1
	
	if curr_sysid != prev_sysid:
		diff_flag = True
		num = 1
		#print('curr=>' + curr_sysid + "   prev=>" + prev_sysid)

	for j in range(df.shape[-1] - 3):
		t.cell(i+1,j).text = str(df.values[i,j])

		#系統代號轉換為中文敘述
		if j == 0:
			if str(df.values[i,j]) == 'BMG':
				t.cell(i+1,j).text = 'B\nM\nG\n煉\n鋼\n鋼\n胚\n'
			elif str(df.values[i,j]) == 'CEO':
				t.cell(i+1,j).text = 'C\nE\nO\n總\n經\n理\n室\n網\n頁\n'
			elif str(df.values[i,j]) == 'CPS':
				t.cell(i+1,j).text = 'C\nP\nS\n冷\n軋\n排\n程\n一\n廠\n'
			elif str(df.values[i,j]) == 'CPS2':
				t.cell(i+1,j).text = 'C\nP\nS\n冷\n軋\n排\n程\n二\n廠\n'
			elif str(df.values[i,j]) == 'CPS3':
				t.cell(i+1,j).text = 'C\nP\nS\n冷\n軋\n排\n程\n三\n廠\n'
			elif str(df.values[i,j]) == 'CPV':
				t.cell(i+1,j).text = 'C\nP\nV\n訂\n單\n接\n收\n'
			elif str(df.values[i,j]) == 'CRM':
				t.cell(i+1,j).text = 'C\nR\nM\n冷\n軋\n傳\n輸\n一\n廠\n'
			elif str(df.values[i,j]) == 'CRM2':
				t.cell(i+1,j).text = 'C\nR\nM\n2\n冷\n軋\n傳\n輸\n二\n廠\n'
			elif str(df.values[i,j]) == 'CRM3':
				t.cell(i+1,j).text = 'C\nR\nM\n3\n冷\n軋\n傳\n輸\n三\n廠\n'
			elif str(df.values[i,j]) == 'CRMW':
				t.cell(i+1,j).text = 'C\nR\nM\nW\n冷\n軋\n生\n產\n管\n理\n(\nW\ne\nb\n)\n'
			elif str(df.values[i,j]) == 'EIS':
				t.cell(i+1,j).text = 'E\nI\nS\n網\n站\n'
			elif str(df.values[i,j]) == 'EQP':
				t.cell(i+1,j).text = 'E\nQ\nP\n設\n備\n管\n理\n系\n統\n'
			elif str(df.values[i,j]) == 'G1':
				t.cell(i+1,j).text = 'G\n1\n生\n計\n部\n網\n頁\n'
			elif str(df.values[i,j]) == 'GNO':
				t.cell(i+1,j).text = 'G\nN\nO\nG\nE\nN\nE\nR\nO\n'
			elif str(df.values[i,j]) == 'HPC':
				t.cell(i+1,j).text = 'H\nP\nC\n熱\n軋\n鋼\n胚\n'
			elif str(df.values[i,j]) == 'HPS':
				t.cell(i+1,j).text = 'H\nP\nS\n熱\n軋\n排\n程\n'
			elif str(df.values[i,j]) == 'HSM':
				t.cell(i+1,j).text = 'H\nS\nM\n熱\n軋\n傳\n輸\n'
			elif str(df.values[i,j]) == 'HSMW':
				t.cell(i+1,j).text = 'H\nS\nM\nW\n熱\n軋\n生\n產\n管\n理\n(\nW\ne\nb\n)\n'
			elif str(df.values[i,j]) == 'ISO':
				t.cell(i+1,j).text = 'I\nS\nO\n規\n章\n管\n理\n'
			elif str(df.values[i,j]) == 'LD':
				t.cell(i+1,j).text = 'L\nD\n聯\n德\n資\n訊\n管\n理\n系\n統\n'
			elif str(df.values[i,j]) == 'MGR':
				t.cell(i+1,j).text = 'M\nG\nR\n主\n管\n查\n詢\n'
			elif str(df.values[i,j]) == 'MIC':
				t.cell(i+1,j).text = 'M\nI\nC\n冶\n金\n規\n範\n'
			elif str(df.values[i,j]) == 'MPS':
				t.cell(i+1,j).text = 'M\nP\nS\n煉\n鋼\n排\n程\n'
			elif str(df.values[i,j]) == 'OEM':
				t.cell(i+1,j).text = 'O\nE\nM\n委\n外\n託\n工\n'
			elif str(df.values[i,j]) == 'ORA':
				t.cell(i+1,j).text = 'O\nR\nA\n\nO\nR\nA\nC\nL\nE\n轉\n檔\n系\n統\n'
			elif str(df.values[i,j]) == 'OTH1':
				t.cell(i+1,j).text = 'O\nT\nH\n1\n其\n它\n(\n大\n電\n腦\n)\n'
			elif str(df.values[i,j]) == 'PCM':
				t.cell(i+1,j).text = 'P\nC\nM\n生\n產\n管\n理\n'
			elif str(df.values[i,j]) == 'PKL':
				t.cell(i+1,j).text = 'P\nK\nL\n包\n裝\n線\n管\n理\n'
			elif str(df.values[i,j]) == 'RTP':
				t.cell(i+1,j).text = 'R\nT\nP\n資\n源\n回\n收\n廠\n生\n產\n管\n理\n網\n'
			elif str(df.values[i,j]) == 'SMP':
				t.cell(i+1,j).text = 'S\nM\nP\n煉\n鋼\n傳\n輸\n系\n統\n'
			elif str(df.values[i,j]) == 'SMPW':
				t.cell(i+1,j).text = 'S\nM\nP\nW\n煉\n鋼\n生\n產\n管\n理\n(\nW\ne\nb\n)\n'
			elif str(df.values[i,j]) == 'STA':
				t.cell(i+1,j).text = 'S\nT\nA\n統\n計\n分\n析\n'
			elif str(df.values[i,j]) == 'T1':
				t.cell(i+1,j).text = 'T\n1\n技\n術\n部\n管\n理\n網\n'
			elif str(df.values[i,j]) == 'T1W':
				t.cell(i+1,j).text = 'T\n1\nW\n各\n廠\n製\n程\n管\n制\n網\n頁\n'
			elif str(df.values[i,j]) == 'TQ1':
				t.cell(i+1,j).text = 'T\nQ\n1\n追\n蹤\n傳\n輸\n'
			elif str(df.values[i,j]) == 'TQ2':
				t.cell(i+1,j).text = 'T\nQ\n2\n實\n驗\n室\n'
			elif str(df.values[i,j]) == 'TQ3':
				t.cell(i+1,j).text = 'T\nQ\n3\n冷\n軋\n品\n管\n'
			elif str(df.values[i,j]) == 'TQ4':
				t.cell(i+1,j).text = 'T\nQ\n4\n煉\n鋼\n品\n管\n'
			elif str(df.values[i,j]) == 'TQ5':
				t.cell(i+1,j).text = 'T\nQ\n5\n熱\n軋\n品\n管\n'
			elif str(df.values[i,j]) == 'TQ6':
				t.cell(i+1,j).text = 'T\nQ\n6\n品\n證\n書\n管\n理\n'
			elif str(df.values[i,j]) == 'TQ7':
				t.cell(i+1,j).text = 'T\nQ\n7\n放\n行\n系\n統\n'
			elif str(df.values[i,j]) == 'TQI':
				t.cell(i+1,j).text = 'T\nQ\nI\n品\n質\n指\n標\n'
			elif str(df.values[i,j]) == 'WIP':
				t.cell(i+1,j).text = 'W\nI\nP\n冷\n軋\n現\n場\n管\n理\n'
			elif str(df.values[i,j]) == 'WIPW':
				t.cell(i+1,j).text = 'W\nI\nP\nW\n當\n日\nW\nI\nP\n'
			elif str(df.values[i,j]) == 'YRD':
				t.cell(i+1,j).text = 'Y\nR\nD\n熱\n軋\n儲\n區\n'
			else:
				t.cell(i+1,j).text = str(df.values[i,j])
				print('"' + str(df.values[i,j]) + '",系統代碼沒有對應中文敘述.\n')

		if j == 0 and diff_flag == False:
			t.cell(i+1,j).text = ''

		#系統功能，前綴序號、最後增加換行符號
		if j == 1:
			t.cell(i+1,j).text = str(num) + '. ' + str(df.values[i,j]) + '\n'

		#上線日期，日期格式轉換
		if j == 3:
			t.cell(i+1,j).text = parser.parse(str(df.values[i,j])).strftime("%Y/%m/%d")

		if diff_flag == False:
			#儲存格合併
			a = t.cell(i, 0)
			b = t.cell(i+1, 0)
			a.merge(b)

	prev_sysid = curr_sysid
	num += 1

# save the doc
document.save('./mis_result_annual.docx')

print("End of prog.")